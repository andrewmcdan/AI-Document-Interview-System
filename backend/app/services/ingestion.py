from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from pathlib import Path
import io
import math
import re
import uuid
from typing import Iterable

import fitz
import pdfplumber
from PIL import Image
import pytesseract
from docx import Document as DocxDocument
from qdrant_client.models import PointStruct
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import update
import tiktoken

from app.core.config import Settings
from app.core.logging import get_logger
from app.db import models
from app.schemas.document import Document as DocumentSchema
from app.schemas.document import DocumentChunk
from app.services.openai_client import OpenAIClient
from app.storage.object_store import ObjectStore
from app.storage.vector_store import VectorStore

log = get_logger()


@dataclass
class IngestionPipeline:
    settings: Settings
    object_store: ObjectStore
    vector_store: VectorStore
    openai_client: OpenAIClient

    async def ingest(
        self,
        file_path: Path,
        document_id: str,
        db: AsyncSession,
        *,
        title: str,
        description: str | None = None,
        owner_id: str | None = None,
        job_id: str | None = None,
    ) -> DocumentSchema:
        """Ingest a document file into storage, DB metadata, and the vector store."""
        object_key = f"{document_id}/{file_path.name}"
        await self._mark_job_running(db, job_id)

        self.object_store.upload_file(file_path, object_key=object_key)

        segments = self.extract_text(file_path)
        chunks = self.chunk_text(segments, document_id=document_id)
        embeddings = self.embed_chunks(chunks)

        try:
            document = await self._persist_document(
                db=db,
                document_id=document_id,
                title=title,
                description=description,
                owner_id=owner_id,
                storage_key=object_key,
            )
            await self.persist_chunks(
                db=db,
                document_id=document_id,
                document_title=document.title,
                owner_id=document.owner_id,
                chunks=chunks,
                embeddings=embeddings,
            )
            await db.commit()
            await db.refresh(document)
            await self._mark_job_complete(db, job_id)
            log.info(
                "ingestion.completed",
                document_id=document_id,
                owner_id=owner_id,
                chunks=len(chunks),
                job_id=job_id,
            )
        except Exception:
            await db.rollback()
            await self._mark_job_failed(db, job_id)
            log.exception("ingestion.failed", document_id=document_id, job_id=job_id)
            raise

        return DocumentSchema(
            id=document.id,
            title=document.title,
            description=document.description,
            owner_id=document.owner_id,
            storage_key=document.storage_key,
            created_at=document.created_at,
        )

    def extract_text(self, file_path: Path) -> list[tuple[int | None, str]]:
        """Return a list of (page_number, text) tuples."""
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            segments = self._extract_pdf(file_path)
        elif suffix == ".docx":
            segments = self._extract_docx(file_path)
        elif suffix in {".txt", ".md"}:
            segments = self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type for ingestion: {file_path.suffix}")

        cleaned: list[tuple[int | None, str]] = []
        for page_no, text in segments:
            normalized = self._normalize_text(text)
            if normalized:
                cleaned.append((page_no, normalized))

        cleaned = self._strip_repeated_headers_footers(cleaned)

        if not cleaned:
            raise ValueError("No text extracted from document; check if the file is scanned or empty.")
        return cleaned

    def chunk_text(self, segments: list[tuple[int | None, str]], document_id: str) -> list[DocumentChunk]:
        """Chunk text using token-aware splitting with overlap."""
        chunk_size = max(1, self.settings.chunk_size_tokens)
        overlap = max(0, min(self.settings.chunk_overlap_tokens, chunk_size - 1))
        encoding = self._encoding()

        chunks: list[DocumentChunk] = []
        chunk_index = 0

        for page_no, text in segments:
            tokens = encoding.encode(text)
            start = 0
            while start < len(tokens):
                end = min(len(tokens), start + chunk_size)
                token_slice = tokens[start:end]
                chunk_text = encoding.decode(token_slice).strip()

                chunk_id = str(uuid.uuid4())
                chunks.append(
                    DocumentChunk(
                        id=chunk_id,
                        document_id=document_id,
                        text=chunk_text,
                        meta={
                            "chunk_index": chunk_index,
                            "page": page_no,
                            "start_token": start,
                            "end_token": end,
                            "text_snippet": chunk_text[:500],
                        },
                    )
                )

                chunk_index += 1
                if end >= len(tokens):
                    break
                start = max(0, end - overlap)

        return chunks

    def embed_chunks(self, chunks: Iterable[DocumentChunk]) -> list[list[float]]:
        if not getattr(self.openai_client.client, "api_key", None):
            raise ValueError("OpenAI API key is missing; set AIDOC_OPENAI_API_KEY to enable embeddings.")
        return [self.openai_client.embed(chunk.text) for chunk in chunks]

    async def persist_chunks(
        self,
        *,
        db: AsyncSession,
        document_id: str,
        document_title: str,
        owner_id: str | None,
        chunks: Iterable[DocumentChunk],
        embeddings: list[list[float]],
    ) -> None:
        if not embeddings:
            raise ValueError("No embeddings produced for document; aborting persistence.")

        self.vector_store.ensure_collection(vector_size=len(embeddings[0]))

        points: list[PointStruct] = []
        chunk_rows: list[models.DocumentChunk] = []

        chunk_list = list(chunks)
        if len(chunk_list) != len(embeddings):
            raise ValueError("Mismatch between chunks and embeddings lengths.")

        for chunk, embedding in zip(chunk_list, embeddings):
            points.append(
                PointStruct(
                    id=chunk.id,
                    vector=embedding,
                    payload={
                        "document_id": document_id,
                        "document_title": document_title,
                        "owner_id": owner_id,
                        "chunk_id": chunk.id,
                        "text": chunk.text,
                        "meta": chunk.meta,
                    },
                )
            )
            chunk_rows.append(
                models.DocumentChunk(
                    id=chunk.id,
                    document_id=document_id,
                    text=chunk.text,
                    meta=chunk.meta,
                )
            )

        self.vector_store.upsert_chunks(points)
        db.add_all(chunk_rows)

    async def _persist_document(
        self,
        *,
        db: AsyncSession,
        document_id: str,
        title: str,
        description: str | None,
        owner_id: str | None,
        storage_key: str,
    ) -> models.Document:
        document = models.Document(
            id=document_id,
            title=title,
            description=description,
            owner_id=owner_id,
            storage_key=storage_key,
        )
        db.add(document)
        await db.flush()
        return document

    async def _mark_job_running(self, db: AsyncSession, job_id: str | None) -> None:
        if not job_id:
            return
        await db.execute(
            update(models.IngestionJob)
            .where(models.IngestionJob.id == job_id)
            .values(status="running", started_at=datetime.utcnow(), error=None)
        )
        await db.commit()

    async def _mark_job_complete(self, db: AsyncSession, job_id: str | None) -> None:
        if not job_id:
            return
        await db.execute(
            update(models.IngestionJob)
            .where(models.IngestionJob.id == job_id)
            .values(status="completed", finished_at=datetime.utcnow(), error=None)
        )
        await db.commit()

    async def _mark_job_failed(self, db: AsyncSession, job_id: str | None) -> None:
        if not job_id:
            return
        await db.execute(
            update(models.IngestionJob)
            .where(models.IngestionJob.id == job_id)
            .values(status="failed", finished_at=datetime.utcnow())
        )
        await db.commit()

    def _extract_pdf(self, file_path: Path) -> list[tuple[int, str]]:
        text_parts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    text_parts.append(page_text)
        if text_parts:
            return list(enumerate(text_parts, start=1))

        # Fallback to PyMuPDF text extraction if pdfplumber finds nothing.
        with fitz.open(file_path) as doc:
            extracted = [(index + 1, page.get_text("text")) for index, page in enumerate(doc)]
            non_empty = [(page_no, text) for page_no, text in extracted if text.strip()]
            if non_empty:
                return non_empty

        # OCR as last resort for scanned PDFs.
        ocr_results = self._ocr_pdf(file_path)
        if ocr_results:
            return ocr_results
        return []

    def _extract_docx(self, file_path: Path) -> list[tuple[int, str]]:
        document = DocxDocument(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        text = "\n".join(paragraphs).strip()
        return [(1, text)] if text else []

    def _extract_txt(self, file_path: Path) -> list[tuple[int, str]]:
        content = file_path.read_text(encoding="utf-8", errors="ignore")
        return [(1, content)]

    def _normalize_text(self, text: str) -> str:
        normalized = re.sub(r"\r\n?", "\n", text)
        normalized = re.sub(r"[ \t]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()

    def _encoding(self):
        try:
            return tiktoken.get_encoding("cl100k_base")
        except Exception:
            return tiktoken.get_encoding("p50k_base")

    def _ocr_pdf(self, file_path: Path) -> list[tuple[int, str]]:
        """Perform OCR on a PDF using Tesseract as a best-effort fallback."""
        try:
            pytesseract.get_tesseract_version()
        except Exception:
            # Tesseract not available; return empty to signal no OCR.
            return []

        results: list[tuple[int, str]] = []
        with fitz.open(file_path) as doc:
            for idx, page in enumerate(doc):
                pix = page.get_pixmap()
                img = Image.open(io.BytesIO(pix.tobytes()))
                text = pytesseract.image_to_string(img)
                if text.strip():
                    results.append((idx + 1, text))
        return results

    def _strip_repeated_headers_footers(self, segments: list[tuple[int | None, str]]) -> list[tuple[int | None, str]]:
        if len(segments) < 2:
            return segments

        first_lines: list[str] = []
        last_lines: list[str] = []
        for _, text in segments:
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            if lines:
                first_lines.append(lines[0])
                last_lines.append(lines[-1])

        def most_common(lines: list[str]) -> str | None:
            if not lines:
                return None
            freq = Counter(lines)
            line, count = freq.most_common(1)[0]
            if count >= max(2, math.ceil(len(segments) * 0.6)):
                return line
            return None

        header = most_common(first_lines)
        footer = most_common(last_lines)

        cleaned_segments: list[tuple[int | None, str]] = []
        for page_no, text in segments:
            lines = [ln.strip() for ln in text.splitlines()]
            if header and lines and lines[0].strip() == header:
                lines = lines[1:]
            if footer and lines and lines[-1].strip() == footer:
                lines = lines[:-1]

            # Remove consecutive duplicate lines
            deduped: list[str] = []
            for ln in lines:
                if deduped and deduped[-1] == ln.strip():
                    continue
                deduped.append(ln.strip())

            combined = "\n".join(ln for ln in deduped if ln.strip()).strip()
            if combined:
                cleaned_segments.append((page_no, combined))

        return cleaned_segments
